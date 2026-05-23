# Import Libraries
from langchain_text_splitters import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    RecursiveJsonSplitter,
)
from tqdm import tqdm
from pypdf import PdfReader
from docx import Document
from openai import OpenAI
import numpy as np
import pandas as pd
import tiktoken
import hashlib
import glob
import json
import csv
import sys
import io
import os

# Redis
import redis
from redis.commands.search.field import TextField, TagField, VectorField
# redis renamed the module from `indexDefinition` to `index_definition` in v5.0+. Try the new name first, then fall back.
try:
    from redis.commands.search.index_definition import IndexDefinition, IndexType
except ImportError:
    from redis.commands.search.indexDefinition import IndexDefinition, IndexType

# Optional PDF→Markdown converter (preserves headings/tables). Falls back to pypdf if unavailable.
try:
    import pymupdf4llm
    _HAS_PYMUPDF4LLM = True
except Exception:
    _HAS_PYMUPDF4LLM = False

# Supress langchain warning messages
import logging
logging.getLogger("langchain.text_splitter").setLevel(logging.ERROR)

# Load environment variables from the PROJECT ROOT .env
from dotenv import load_dotenv
from pathlib import Path
load_dotenv(dotenv_path=Path(__file__).resolve().parent.parent / ".env")

openai_key = os.getenv("OPENAI_API_KEY")
az_connection_str = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
# AWS keys are auto-detected by boto3 from AWS_ACCESS_KEY_ID / AWS_SECRET_ACCESS_KEY — no need to read them here.

# Initialize OpenAI Client
openai_client = OpenAI(api_key=openai_key)

# Initialize Redis Client
redis_client = redis.Redis(host="localhost", port=6379)

# ---------------------------------------------------------------------------
# Parameters
# ---------------------------------------------------------------------------
DOC_PREFIX = "doc:"                         # RediSearch Key Prefix for the Index
EMBEDDING_MODEL = "text-embedding-3-large"  # Embedding model
EMBEDDING_DIMENSIONS = 3072                 # Embedding dimensions
CHUNK_TOKENS = 1200                         # Target chunk size in tokens (~900 words)
CHUNK_OVERLAP_TOKENS = 150                  # Overlap in tokens (~12% of chunk)
CSV_ROWS_PER_CHUNK = 25                     # For CSV/XLSX: rows grouped per chunk

# Vector index settings — env-driven so backend & indexer stay in sync.
INDEX_TYPE = os.getenv("INDEX_TYPE", "HNSW")             # FLAT or HNSW
EF_RUNTIME = int(os.getenv("EF_RUNTIME", "64"))          # HNSW: query-time candidates

# Tokenizer for chunk sizing (text-embedding-3-* uses cl100k_base)
_ENCODER = tiktoken.get_encoding("cl100k_base")


def tiktoken_len(text: str) -> int:
    """Token length using cl100k_base (used by text-embedding-3-*)."""
    if not text:
        return 0
    return len(_ENCODER.encode(text))


def content_hash(text: str) -> str:
    """Short stable hash of chunk content for idempotent doc IDs."""
    return hashlib.sha1(text.encode("utf-8")).hexdigest()[:10]


# ---------------------------------------------------------------------------
# Helper Functions
# ---------------------------------------------------------------------------
def get_embedding(text):
    text = text.replace("\n", " ")
    return openai_client.embeddings.create(input=[text], model=EMBEDDING_MODEL).data[0].embedding


def _recursive_splitter(separators=None):
    """Token-length aware recursive splitter."""
    return RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_TOKENS,
        chunk_overlap=CHUNK_OVERLAP_TOKENS,
        length_function=tiktoken_len,
        separators=separators or ["\n\n", "\n", ". ", " ", ""],
    )


def split_text(doc_text):
    """
    Generic text splitting. Returns a list of dicts:
        [{"content": str, "heading": str}]
    """
    try:
        chunks = _recursive_splitter().split_text(doc_text)
        return [{"content": c, "heading": ""} for c in chunks]
    except Exception as e:
        print(f"Error while splitting text: {str(e)}")
        return []


def split_markdown(md_text):
    """
    Markdown-aware splitter: first split by headings, then recursively size each section.
    Preserves heading metadata as a separate field per chunk.
    """
    try:
        headers_to_split_on = [
            ("#", "h1"),
            ("##", "h2"),
            ("###", "h3"),
            ("####", "h4"),
        ]
        md_splitter = MarkdownHeaderTextSplitter(headers_to_split_on=headers_to_split_on)
        header_chunks = md_splitter.split_text(md_text)

        recursive = _recursive_splitter()
        out = []
        for hc in header_chunks:
            # Build a readable breadcrumb of the headings present on this chunk
            crumbs = []
            for _, key in headers_to_split_on:
                if key in hc.metadata:
                    crumbs.append(str(hc.metadata[key]))
            heading = " > ".join(crumbs)

            text = hc.page_content
            # Prepend heading inside the chunk text too (helps both BM25 and embeddings)
            text_with_heading = f"{heading}\n\n{text}" if heading else text

            for piece in recursive.split_text(text_with_heading):
                out.append({"content": piece, "heading": heading})
        return out
    except Exception as e:
        print(f"Error while splitting markdown: {str(e)}")
        # Fallback to plain text splitter
        return split_text(md_text)


def split_json(doc_json):
    """JSON-aware splitter (size budget is approximate char count for the JSON splitter)."""
    try:
        json_splitter = RecursiveJsonSplitter(max_chunk_size=CHUNK_TOKENS * 4)  # rough chars-per-token mapping
        chunks = json_splitter.split_json(json_data=doc_json)
        return [{"content": json.dumps(c, ensure_ascii=False), "heading": ""} for c in chunks]
    except Exception as e:
        print(f"Error while splitting JSON: {str(e)}")
        return []


def split_tabular_rows(rows, headers):
    """
    Group rows for CSV/XLSX into chunks. Prepend the header row to every chunk so each chunk
    is self-describing. `rows` is a list of dicts; `headers` is a list of column names.
    """
    chunks = []
    if not rows:
        return chunks

    header_line = "Columns: " + ", ".join(headers)
    for i in range(0, len(rows), CSV_ROWS_PER_CHUNK):
        batch = rows[i:i + CSV_ROWS_PER_CHUNK]
        body_lines = []
        for j, row in enumerate(batch, start=i + 1):
            row_str = " | ".join(f"{k}: {v}" for k, v in row.items())
            body_lines.append(f"Row {j}: {row_str}")
        content = header_line + "\n" + "\n".join(body_lines)
        chunks.append({"content": content, "heading": f"rows {i + 1}-{i + len(batch)}"})
    return chunks


# ---------------------------------------------------------------------------
# Index management
# ---------------------------------------------------------------------------
def _schema_fields():
    """Common schema fields used by both FLAT and HNSW indexes."""
    return (
        TextField("index_name"),
        TextField("file_name"),
        TextField("chunk"),
        TextField("content"),          # used for BM25 full-text search
        TextField("source"),
        TextField("heading"),          # heading breadcrumb (markdown/docx/csv-rows)
        TextField("prev_id"),          # neighbor linkage
        TextField("next_id"),
        TagField("tag"),
    )


def create_index(index_name):
    try:
        redis_client.ft(index_name).info()
        print("Index already exists!")
        return
    except Exception:
        pass

    if INDEX_TYPE == "FLAT":
        vector = VectorField(
            "vector",
            "FLAT",
            {
                "TYPE": "FLOAT32",
                "DIM": EMBEDDING_DIMENSIONS,
                "DISTANCE_METRIC": "COSINE",
            },
        )
    elif INDEX_TYPE == "HNSW":
        vector = VectorField(
            "vector",
            "HNSW",
            {
                "TYPE": "FLOAT32",
                "DIM": EMBEDDING_DIMENSIONS,
                "DISTANCE_METRIC": "COSINE",
                "M": 16,
                "EF_CONSTRUCTION": 200,
                "EF_RUNTIME": EF_RUNTIME,
                "EPSILON": 0.01,
            },
        )
    else:
        raise ValueError("Invalid Index Type Passed")

    schema = _schema_fields() + (vector,)
    definition = IndexDefinition(prefix=[DOC_PREFIX], index_type=IndexType.HASH)
    redis_client.ft(index_name).create_index(fields=schema, definition=definition)
    print(f"Index {index_name} created!")


def insert_index(filename, source, doc_chunks, index_name):
    """
    Insert chunks into Redis with content-hash IDs and prev/next chunk linkage.

    `doc_chunks` is a list of dicts: [{"content": str, "heading": str}]
    """
    try:
        if not doc_chunks:
            return

        # Pre-compute unique IDs so we can wire prev/next before HSET
        ids = []
        for i, ch in enumerate(doc_chunks):
            h = content_hash(ch["content"])
            ids.append(f"doc:{index_name}:{filename}:c{i+1}:{h}")

        # Create embeddings + write
        pipe = redis_client.pipeline()
        for i, ch in enumerate(doc_chunks):
            vec = np.array(get_embedding(ch["content"]), dtype=np.float32).tobytes()
            data = {
                "index_name": index_name,
                "file_name": filename,
                "chunk": f"chunk_{i+1}",
                "content": ch["content"],
                "source": source,
                "heading": ch.get("heading", "") or "",
                "prev_id": ids[i - 1] if i > 0 else "",
                "next_id": ids[i + 1] if i < len(ids) - 1 else "",
                "tag": "OpenAIEmbeddings",
                "vector": vec,
            }
            pipe.hset(ids[i], mapping=data)

        pipe.execute()
    except Exception as e:
        print(str(e))


# ---------------------------------------------------------------------------
# File readers
# ---------------------------------------------------------------------------
def read_txt(file_path, local=True):
    doc_chunks = []
    try:
        if local:
            with open(file_path, "r") as f:
                doc_text = f.read()
        else:
            doc_text = file_path.decode("utf-8")

        # If file looks like markdown, use heading-aware splitter
        if local and isinstance(file_path, str) and file_path.lower().endswith(".md"):
            doc_chunks = split_markdown(doc_text)
        else:
            doc_chunks = split_text(doc_text)
    except Exception as e:
        print(f"Error: {str(e)}")
    return doc_chunks


def read_pdf(file_path):
    """
    Prefer pymupdf4llm (PDF → Markdown, preserves structure). Falls back to pypdf plain text.
    """
    doc_chunks = []
    try:
        if _HAS_PYMUPDF4LLM:
            # pymupdf4llm accepts a path string or BytesIO
            md_text = pymupdf4llm.to_markdown(file_path)
            doc_chunks = split_markdown(md_text)
        else:
            reader = PdfReader(file_path)
            doc_text = "\n".join(p.extract_text() or "" for p in reader.pages)
            doc_chunks = split_text(doc_text)
    except Exception as e:
        print(f"Error: {str(e)}")
    return doc_chunks


def read_docx(file_path):
    """
    DOCX: group paragraphs by heading style so each chunk carries its section heading.
    """
    doc_chunks = []
    try:
        doc = Document(file_path)
        sections = []   # list of (heading, [paragraphs])
        current_heading = ""
        current_body = []

        for para in doc.paragraphs:
            style = (para.style.name or "").lower() if para.style else ""
            text = (para.text or "").strip()
            if not text:
                continue
            if style.startswith("heading"):
                # Flush previous section
                if current_body:
                    sections.append((current_heading, current_body))
                current_heading = text
                current_body = []
            else:
                current_body.append(text)
        if current_body:
            sections.append((current_heading, current_body))

        # Fallback: no headings found, treat the whole doc as one section
        if not sections:
            sections = [("", [p.text for p in doc.paragraphs if p.text])]

        recursive = _recursive_splitter()
        for heading, body in sections:
            joined = "\n".join(body).strip()
            if not joined:
                continue
            # Prepend heading to text so the embedding sees the section context
            text_with_heading = f"{heading}\n\n{joined}" if heading else joined
            for piece in recursive.split_text(text_with_heading):
                doc_chunks.append({"content": piece, "heading": heading})
    except Exception as e:
        print(f"Error: {str(e)}")
    return doc_chunks


def read_json(file_path, local=True):
    doc_chunks = []
    try:
        if local:
            with open(file_path, "r") as f:
                doc_text = json.load(f)
        else:
            doc_text = json.loads(file_path)
        doc_chunks = split_json(doc_text)
    except json.JSONDecodeError as e:
        print(f"Error: {str(e)}")
    except Exception as e:
        print(f"Error: {str(e)}")
    return doc_chunks


def read_csv(file_path, local=True):
    doc_chunks = []
    try:
        rows = []
        headers = []
        if local:
            with open(file_path, mode="r", encoding="utf-8") as f:
                reader = csv.DictReader(f)
                headers = reader.fieldnames or []
                for row in reader:
                    rows.append(row)
        else:
            reader = csv.DictReader(file_path)
            headers = reader.fieldnames or []
            for row in reader:
                rows.append(row)
        doc_chunks = split_tabular_rows(rows, headers)
    except Exception as e:
        print(f"Error: {str(e)}")
    return doc_chunks


def read_xlsx(file_path):
    doc_chunks = []
    try:
        df = pd.read_excel(file_path, engine="openpyxl")
        headers = list(df.columns.astype(str))
        rows = df.astype(str).to_dict(orient="records")
        doc_chunks = split_tabular_rows(rows, headers)
    except Exception as e:
        print(f"Error: {str(e)}")
    return doc_chunks


# ---------------------------------------------------------------------------
# Source indexers
# ---------------------------------------------------------------------------
def index_local(src_directory, index_name):
    create_index(index_name)
    files = glob.glob(f"{src_directory}/*.*")
    source = "Local"

    for i in tqdm(range(len(files)), desc="Indexing Files "):
        _, filename = os.path.split(files[i])
        ext = filename.split(".")[-1].lower()
        doc_chunks = []

        if ext in ("txt", "md"):
            doc_chunks = read_txt(files[i])
        elif ext == "pdf":
            doc_chunks = read_pdf(files[i])
        elif ext == "docx":
            doc_chunks = read_docx(files[i])
        elif ext == "json":
            doc_chunks = read_json(files[i])
        elif ext == "csv":
            doc_chunks = read_csv(files[i])
        elif ext == "xlsx":
            doc_chunks = read_xlsx(files[i])

        insert_index(filename, source, doc_chunks, index_name)


def index_azure(container_name, index_name):
    from azure.storage.blob import BlobServiceClient

    blob_service_client = BlobServiceClient.from_connection_string(az_connection_str)
    container_client = blob_service_client.get_container_client(container_name)

    create_index(index_name)
    source = "Azure"
    blobs = list(container_client.list_blobs())

    for i in tqdm(range(len(blobs)), desc="Indexing Files "):
        filename = blobs[i].name
        ext = filename.split(".")[-1].lower()
        blob_client = blob_service_client.get_blob_client(container=container_name, blob=filename)
        file_data = blob_client.download_blob().readall()

        doc_chunks = []
        if ext in ("txt", "md"):
            doc_chunks = read_txt(file_data, False)
        elif ext == "pdf":
            doc_chunks = read_pdf(io.BytesIO(file_data))
        elif ext == "docx":
            doc_chunks = read_docx(io.BytesIO(file_data))
        elif ext == "json":
            doc_chunks = read_json(file_data, False)
        elif ext == "csv":
            csv_data = blob_client.download_blob().content_as_text()
            doc_chunks = read_csv(io.StringIO(csv_data), False)
        elif ext == "xlsx":
            doc_chunks = read_xlsx(io.BytesIO(file_data))

        insert_index(filename, source, doc_chunks, index_name)


def index_aws(bucket_name, index_name):
    import boto3

    # boto3 auto-detects credentials from AWS_ACCESS_KEY_ID / AWS_SECRET_ACCESS_KEY (and AWS_DEFAULT_REGION) in the environment.
    s3_client = boto3.client("s3")

    result = s3_client.list_objects_v2(Bucket=bucket_name)
    if "Contents" not in result:
        print("The bucket is empty.")
        return

    create_index(index_name)
    source = "AWS"

    for i in tqdm(range(len(result["Contents"])), desc="Indexing Files "):
        filename = result["Contents"][i]["Key"]
        ext = filename.split(".")[-1].lower()
        response = s3_client.get_object(Bucket=bucket_name, Key=filename)
        file_data = response["Body"].read()

        doc_chunks = []
        if ext in ("txt", "md"):
            doc_chunks = read_txt(file_data, False)
        elif ext == "pdf":
            doc_chunks = read_pdf(io.BytesIO(file_data))
        elif ext == "docx":
            doc_chunks = read_docx(io.BytesIO(file_data))
        elif ext == "json":
            doc_chunks = read_json(file_data, False)
        elif ext == "csv":
            doc_chunks = read_csv(io.StringIO(file_data.decode("utf-8")), False)
        elif ext == "xlsx":
            doc_chunks = read_xlsx(io.BytesIO(file_data))

        insert_index(filename, source, doc_chunks, index_name)


# ---------------------------------------------------------------------------
# Entrypoint
# ---------------------------------------------------------------------------
if len(sys.argv) != 4:
    print("Invalid number of parameter passed")
    exit()

src = sys.argv[1]
src_dir = sys.argv[2]
index_name = sys.argv[3]

if src == "local":
    print("Indexing local files ...")
    index_local(src_dir, index_name)
elif src == "azure":
    print("Indexing Azure Blob Container files ...")
    index_azure(src_dir, index_name)
elif src == "aws":
    print("Indexing AWS S3 files ...")
    index_aws(src_dir, index_name)
else:
    print("Invalid Source Passed")
